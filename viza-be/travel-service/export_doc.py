from tempfile import NamedTemporaryFile

from docx import Document

from export_summary import (
    TABLE_KEYS,
    build_itinery_rows,
    get_table_headers,
    localize_export_text,
    localize_itinery_rows,
    normalize_export_language,
)


def _section_title(language):
    return "itinerary" if language == "en" else "行程"


def _add_table(doc, rows, language):
    localized_rows = localize_itinery_rows(rows, language)
    headers = get_table_headers(language)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    if localized_rows:
        for row in localized_rows:
            cells = table.add_row().cells
            for index, key in enumerate(TABLE_KEYS):
                cells[index].text = str(row.get(key, "-"))
    else:
        cells = table.add_row().cells
        cells[0].text = "No itinerary" if language == "en" else "暂无行程"
        for index in range(1, len(headers)):
            cells[index].text = "-"


def _add_day_details(doc, itinerary, language):
    if not itinerary:
        doc.add_paragraph("No itinerary generated." if language == "en" else "暂无行程。")
        return

    for day in itinerary:
        day_label = day.get("day", "-")
        city = localize_export_text(day.get("city", "Unknown"), language, "route")
        heading = (
            f"Day {day_label} - {city}"
            if language == "en"
            else f"第 {day_label} 天 - {city}"
        )
        doc.add_heading(heading, level=1)

        doc.add_paragraph("Activities:" if language == "en" else "活动：")
        for activity in day.get("activities", []):
            doc.add_paragraph(f"- {localize_export_text(activity, language, 'name')}")

        doc.add_paragraph("Dining:" if language == "en" else "餐饮：")
        for food in day.get("food", []):
            doc.add_paragraph(f"- {localize_export_text(food, language, 'name')}")

        cost_label = "Cost" if language == "en" else "预算"
        doc.add_paragraph(
            f"{cost_label}: {localize_export_text(day.get('cost', 'N/A'), language, 'details')}"
        )


def _add_language_section(doc, itinerary, rows, language):
    doc.add_heading(_section_title(language), 0)
    _add_table(doc, rows, language)
    _add_day_details(doc, itinerary, language)


def export_to_word(itinerary, state=None):
    state = state or {}
    export_language = normalize_export_language(state.get("export_language"))
    doc = Document()
    rows = build_itinery_rows(itinerary, state)

    if export_language == "bilingual":
        _add_language_section(doc, itinerary, rows, "zh")
        doc.add_page_break()
        _add_language_section(doc, itinerary, rows, "en")
    else:
        _add_language_section(doc, itinerary, rows, export_language)

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        file_path = tmp.name

    doc.save(file_path)
    return file_path
